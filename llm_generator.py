# llm_generator.py
from __future__ import annotations

import random
import re
from typing import Any, Dict, List, Optional, Tuple

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

from llm_clients import BaseLLMClient, LLMResult


# ------------------------------
# Utilities
# ------------------------------

def clean_text(text: str) -> str:
    """Remove illegal XML/control chars for Word compatibility."""
    if not text:
        return ""
    return re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", str(text))


# ------------------------------
# Exporters
# ------------------------------

class TxtExporter:
    def export(self, questions: List[Dict[str, Any]], filename: str = "assessment_questions.txt",
               title: str = "AI-Generated Quantitative Math Assessment") -> str:
        chunks: List[str] = []
        for i, q in enumerate(questions, 1):
            chunks.append(self._format_question(q, i, title))
        content = "\n".join(chunks)
        with open(filename, "w", encoding="utf-8") as f:
            f.write(content)
        return filename

    def _format_question(self, q: Dict[str, Any], number: int, title: str) -> str:
        out: List[str] = []
        if number == 1:
            out.append(f"@title {title}")
            out.append("@description Comprehensive math assessment covering various curriculum topics")
            out.append("")
        out.append(f"@question {clean_text(q['question'])}")
        out.append("@instruction Choose the correct option")
        out.append(f"@difficulty {clean_text(q['difficulty'])}")
        out.append(f"@Order {number}")
        for i, opt in enumerate(q["options"]):
            tag = "@@option" if i == q["correct_index"] else "@option"
            out.append(f"{tag} {clean_text(opt)}")
        out.append(f"@explanation {clean_text(q['explanation'])}")
        out.append(f"@subject {clean_text(q['subject'])}")
        out.append(f"@unit {clean_text(q['unit'])}")
        out.append(f"@topic {clean_text(q['topic'])}")
        out.append("@plusmarks 1")
        out.append("")
        return "\n".join(out)


class WordExporter:
    """
    Clean, modular DOCX exporter with safe text handling and optional image embedding.
    """

    def __init__(self, filename: str = "math_assessment.docx"):
        self.filename = filename
        self.doc = Document()

    def add_title(self) -> None:
        self.doc.add_heading("AI-Generated Math Assessment Questions", 0)
        self.doc.add_paragraph(
            "This document contains AI-generated math assessment questions "
            "following the specified curriculum and format requirements."
        )

    def add_question(self, q: Dict[str, Any], number: int, image_path: Optional[str] = None) -> None:
        self._add_question_block(q, number)
        self._add_options(q["options"], q["correct_index"])
        self._add_details(q)
        self._add_explanation(q["explanation"])
        if image_path:
            self.doc.add_picture(image_path, width=Inches(4))
        if number > 0:
            self.doc.add_page_break()

    def _add_question_block(self, q: Dict[str, Any], number: int) -> None:
        self.doc.add_heading(f"Question {number}", level=1)
        p = self.doc.add_paragraph()
        p.add_run("Question: ").bold = True
        p.add_run(clean_text(q["question"]))

    def _add_options(self, options: List[str], correct_index: int) -> None:
        self.doc.add_paragraph("Options:", style="Heading 2")
        for idx, opt in enumerate(options):
            letter = chr(65 + idx)
            p = self.doc.add_paragraph()
            if idx == correct_index:
                p.add_run(f"({letter}) {clean_text(opt)} ").bold = True
                p.add_run("✓").bold = True
            else:
                p.add_run(f"({letter}) {clean_text(opt)}")

    def _add_details(self, q: Dict[str, Any]) -> None:
        self.doc.add_paragraph("Assessment Details:", style="Heading 2")
        details = self.doc.add_paragraph()
        details.add_run(f"Difficulty: {clean_text(q['difficulty'])}\n")
        details.add_run(f"Subject: {clean_text(q['subject'])}\n")
        details.add_run(f"Unit: {clean_text(q['unit'])}\n")
        details.add_run(f"Topic: {clean_text(q['topic'])}\n")

    def _add_explanation(self, explanation: str) -> None:
        p = self.doc.add_paragraph()
        p.add_run("Explanation: ").bold = True
        p.add_run(clean_text(explanation))

    def save(self) -> str:
        self.doc.save(self.filename)
        return self.filename


# ------------------------------
# Core Generator
# ------------------------------

class AIQuestionGenerator:
    def __init__(self, llm_client: Optional[BaseLLMClient] = None) -> None:
        self.llm = llm_client

    # ---------- Public API ----------
    def generate_assessment(self, n: int = 2) -> List[Dict[str, Any]]:
        topic_pool = [
            ("Circles (Area, circumference)", "moderate"),
            ("Quadratic Equations & Functions (Finding roots/solutions, graphing)", "moderate"),
            ("Coordinate Geometry", "easy"),
            ("Fractions, Decimals, & Percents", "moderate"),
            ("Interpreting Variables", "easy"),
        ]
        picks = random.sample(topic_pool, k=min(n, len(topic_pool)))
        while len(picks) < n:  # if n > pool size
            picks.append(random.choice(topic_pool))

        questions: List[Dict[str, Any]] = []
        for topic, diff in picks:
            q = self.generate_question(topic, diff)
            q["options"] = self._ensure_five(q["options"])
            if not (0 <= int(q["correct_index"]) < 5):
                q["correct_index"] = 0
            questions.append(q)
        return questions

    def generate_question(self, topic: str, difficulty: str = "moderate") -> Dict[str, Any]:
        # Prefer LLM; fallback if it fails or returns invalid
        if self.llm:
            prompt = self._build_prompt(topic, difficulty)
            result: LLMResult = self.llm.generate(prompt)
            if result.ok and result.data:
                return self._validate_and_fix_llm(result.data, topic, difficulty)
        return self._fallback_by_topic(topic, difficulty)

    # ---------- Prompting ----------
    def _build_prompt(self, topic: str, difficulty: str) -> str:
        schema = {
            "question": "string (use LaTeX if needed)",
            "options": ["string", "string", "string", "string", "string"],
            "correct_index": "integer (0-4)",
            "explanation": "string (step-by-step; LaTeX allowed)",
        }
        return (
            "Generate ONE multiple-choice math question as JSON ONLY (no extra text).\n"
            f"Topic: {topic}\n"
            f"Difficulty: {difficulty}\n"
            "Requirements:\n"
            "- Use LaTeX for math (e.g., $x^2$, \\frac{a}{b}).\n"
            "- Provide EXACTLY 5 unique options in an array.\n"
            "- Set correct_index to the correct option's index (0-4).\n"
            "- Keys must be: question, options, correct_index, explanation.\n\n"
            f"Return a compact JSON matching this schema: {schema}\n"
        )

    # ---------- Validation / Repair ----------
    def _validate_and_fix_llm(self, data: Dict[str, Any], topic: str, difficulty: str) -> Dict[str, Any]:
        q_txt = clean_text(str(data.get("question", "")).strip())
        options_in = data.get("options", [])
        try:
            correct_index = int(data.get("correct_index", -1))
        except Exception:
            correct_index = -1
        explanation = clean_text(str(data.get("explanation", "")).strip())

        options = self._ensure_five([clean_text(o) for o in options_in])

        if not (0 <= correct_index < 5):
            correct_index = 0
        if not q_txt or not explanation:
            return self._fallback_by_topic(topic, difficulty)

        subject, unit, mapped_topic = self._map_curriculum(topic)
        return {
            "question": q_txt,
            "options": options,
            "correct_index": correct_index,
            "explanation": explanation,
            "subject": subject,
            "unit": unit,
            "topic": mapped_topic,
            "difficulty": difficulty,
            "has_image": True if "coordinate" in mapped_topic.lower() else False,
        }

    # ---------- Curriculum mapping ----------
    def _map_curriculum(self, topic: str) -> Tuple[str, str, str]:
        t = topic.lower()
        if "circle" in t:
            return ("Quantitative Math", "Geometry and Measurement", "Circles (Area, circumference)")
        if "quadratic" in t:
            return ("Quantitative Math", "Algebra",
                    "Quadratic Equations & Functions (Finding roots/solutions, graphing)")
        if "coordinate" in t:
            return ("Quantitative Math", "Geometry and Measurement", "Coordinate Geometry")
        if "fraction" in t or "percent" in t:
            return ("Quantitative Math", "Numbers and Operations", "Fractions, Decimals, & Percents")
        if "interpreting variables" in t or "linear" in t:
            return ("Quantitative Math", "Algebra", "Interpreting Variables")
        return ("Quantitative Math", "Problem Solving", "Algebra")

    # ---------- Fallback generators ----------
    def _fallback_by_topic(self, topic: str, difficulty: str) -> Dict[str, Any]:
        t = topic.lower()
        if "circle" in t:
            return self._gen_circle(difficulty)
        if "quadratic" in t:
            return self._gen_quadratic(difficulty)
        if "coordinate" in t:
            return self._gen_coordinate(difficulty)
        if "fraction" in t or "percent" in t:
            return self._gen_fraction(difficulty)
        return self._gen_linear(difficulty)

    def _ensure_five(self, options: List[str]) -> List[str]:
        uniq: List[str] = []
        seen = set()
        for o in options:
            s = (o or "").strip()
            if s and s not in seen:
                uniq.append(s)
                seen.add(s)
            if len(uniq) == 5:
                break
        while len(uniq) < 5:
            filler = str(random.randint(1, 99))
            if filler not in seen:
                uniq.append(filler)
                seen.add(filler)
        return uniq[:5]

    def _gen_circle(self, difficulty: str) -> Dict[str, Any]:
        r = random.randint(3, 12)
        area = r * r
        correct = f"${area}\\pi$"
        options = self._ensure_five([
            correct,
            f"${2*r}\\pi$",
            f"${r}\\pi$",
            f"${2*area}\\pi$",
            f"${max(1, area//2)}\\pi$",
        ])
        idx = options.index(correct) if correct in options else 0
        return {
            "question": f"A circle has a radius of {r} units. What is the area of the circle?",
            "options": options,
            "correct_index": idx,
            "explanation": f"Area formula: $A=\\pi r^2$. With $r={r}$, $A=\\pi\\cdot {r}^2={area}\\pi$ square units.",
            "subject": "Quantitative Math",
            "unit": "Geometry and Measurement",
            "topic": "Circles (Area, circumference)",
            "difficulty": difficulty,
            "has_image": False,
        }

    def _gen_quadratic(self, difficulty: str) -> Dict[str, Any]:
        a, b = random.sample(range(-6, 7), 2)
        while b == a:
            a, b = random.sample(range(-6, 7), 2)
        B = -(a + b)
        C = a * b
        eq = f"$x^2 {'+' if B>=0 else '-'} {abs(B)}x {'+' if C>=0 else '-'} {abs(C)}=0$"
        correct = f"$x={a}$ and $x={b}$"
        options = self._ensure_five([
            correct,
            f"$x={a+1}$ and $x={b+1}$",
            f"$x={-a}$ and $x={-b}$",
            f"$x={a}$ and $x={b+2}$",
            f"$x={a-1}$ and $x={b}$",
        ])
        idx = options.index(correct) if correct in options else 0
        return {
            "question": f"If {eq}, what are all possible values of $x$?",
            "options": options,
            "correct_index": idx,
            "explanation": f"Factor: $(x-{a})(x-{b})=0$ so $x={a}$ or $x={b}$.",
            "subject": "Quantitative Math",
            "unit": "Algebra",
            "topic": "Quadratic Equations & Functions (Finding roots/solutions, graphing)",
            "difficulty": difficulty,
            "has_image": False,
        }

    def _gen_coordinate(self, difficulty: str) -> Dict[str, Any]:
        target_x = random.randint(-5, 5)
        target_y = random.randint(-5, 5)
        pts: List[Tuple[int, int]] = [(target_x, target_y)]
        xs = {target_x}
        while len(pts) < 5:
            x = random.randint(-5, 5)
            if x in xs:
                continue
            y = random.randint(-5, 5)
            pts.append((x, y))
            xs.add(x)
        random.shuffle(pts)
        letters = ["A", "B", "C", "D", "E"]
        idx = pts.index((target_x, target_y))
        answer_letter = letters[idx]
        return {
            "question": f"Which point on the coordinate plane has an $x$-coordinate of {target_x}?",
            "options": letters.copy(),
            "correct_index": idx,
            "explanation": f"Point {answer_letter} is at $({target_x}, {target_y})$ so its $x$-coordinate is {target_x}.",
            "subject": "Quantitative Math",
            "unit": "Geometry and Measurement",
            "topic": "Coordinate Geometry",
            "difficulty": difficulty,
            "has_image": True,
            "points_data": {letters[i]: pts[i] for i in range(5)},
        }

    def _gen_fraction(self, difficulty: str) -> Dict[str, Any]:
        total = random.randint(40, 120)
        percentage = random.choice([25, 30, 40, 50, 60, 75, 80])
        category = random.choice(["wearing glasses", "playing sports", "taking music lessons"])
        correct = (total * percentage) // 100
        wrongs = list({correct + d for d in [-10, -5, 5, 10] if correct + d > 0})
        options = self._ensure_five([str(correct)] + [str(w) for w in wrongs])
        random.shuffle(options)
        idx = options.index(str(correct))
        return {
            "question": f"In a group of {total} students, {percentage}% are {category}. How many students are {category}?",
            "options": options,
            "correct_index": idx,
            "explanation": f"Compute {percentage}% of {total}: $\\frac{{{percentage}}}{{100}}\\times {total}={correct}$.",
            "subject": "Quantitative Math",
            "unit": "Numbers and Operations",
            "topic": "Fractions, Decimals, & Percents",
            "difficulty": difficulty,
            "has_image": False,
        }

    def _gen_linear(self, difficulty: str) -> Dict[str, Any]:
        a = random.randint(2, 10)
        c = random.randint(1, a - 1)
        d = random.randint(5, 20)
        x_target = random.randint(2, 5)
        b = d - x_target * (a - c)
        expr = f"${a}x {'+' if b>=0 else '-'} {abs(b)} = {c}x + {d}$"
        wrongs = [x_target + 1, max(1, x_target - 1), x_target + 2, x_target * 2]
        options = self._ensure_five([str(x_target)] + [str(w) for w in wrongs])
        random.shuffle(options)
        idx = options.index(str(x_target))
        return {
            "question": f"If {expr}, what is the value of $x$?",
            "options": options,
            "correct_index": idx,
            "explanation": f"${a}x-{c}x={d}-{b}$ so ${(a-c)}x={d-b}$, hence $x={x_target}$.",
            "subject": "Quantitative Math",
            "unit": "Algebra",
            "topic": "Interpreting Variables",
            "difficulty": difficulty,
            "has_image": False,
        }

    # ---------- Images ----------
    def create_graph_image(self, question_data: Dict[str, Any], filename: str) -> Optional[str]:
        if not question_data.get("has_image") or "points_data" not in question_data:
            return None
        fig, ax = plt.subplots(figsize=(6, 6))
        ax.grid(True, alpha=0.3)
        ax.set_xlim(-6, 6)
        ax.set_ylim(-6, 6)
        ax.axhline(y=0, color="k", linewidth=0.7)
        ax.axvline(x=0, color="k", linewidth=0.7)
        for label, (x, y) in question_data["points_data"].items():
            ax.plot(x, y, marker="o")
            ax.annotate(label, (x, y), xytext=(6, 6), textcoords="offset points")
        ax.set_xlabel("x")
        ax.set_ylabel("y")
        ax.set_title("Coordinate Plane")
        fig.tight_layout()
        fig.savefig(filename, dpi=150, bbox_inches="tight")
        plt.close(fig)
        return filename
